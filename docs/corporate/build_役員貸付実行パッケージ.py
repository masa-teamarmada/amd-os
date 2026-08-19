from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "/Users/masa/projects/AMD/amd-os/docs/corporate/20260816_役員貸付実行パッケージ.docx"
doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)

styles = doc.styles
for name, size, color in [("Normal", 11, None), ("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5")]:
    st = styles[name]
    st.font.name = "Hiragino Sans"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Hiragino Sans")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Hiragino Sans")
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans")
    st.font.size = Pt(size)
    if color: st.font.color.rgb = RGBColor.from_string(color)

def para(text="", bold=False, size=None, align=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align: p.alignment = align
    r = p.add_run(text); r.bold = bold; r.font.name = "Hiragino Sans"; r._element.rPr.rFonts.set(qn("w:ascii"), "Hiragino Sans"); r._element.rPr.rFonts.set(qn("w:hAnsi"), "Hiragino Sans"); r._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans")
    if size: r.font.size = Pt(size)
    return p

def heading(text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)

def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.autofit = False
    for i, h in enumerate(headers):
        cell=t.rows[0].cells[i]; cell.width=Inches(widths[i]); cell.text=h
        for r in cell.paragraphs[0].runs: r.bold=True
    for row in rows:
        cells=t.add_row().cells
        for i, value in enumerate(row):
            cells[i].width=Inches(widths[i]); cells[i].text=str(value)
    for row in t.rows:
        for cell in row.cells:
            tcPr=cell._tc.get_or_add_tcPr(); mar=OxmlElement("w:tcMar")
            for side in ("top","start","bottom","end"):
                n=OxmlElement(f"w:{side}"); n.set(qn("w:w"), "100"); n.set(qn("w:type"), "dxa"); mar.append(n)
            tcPr.append(mar)
    return t

para("役員貸付実行パッケージ", True, 22, WD_ALIGN_PARAGRAPH.CENTER, 4)
para("株式会社チームアルマダ / 2026年8月16日作成 / 署名前ドラフト", False, 11, WD_ALIGN_PARAGRAPH.CENTER, 18)
heading("1. 確定前提")
for x in ["株式会社チームアルマダから山地正洋へ2,000,000円を貸し付ける。", "PayPay銀行の新規500万円融資が株式会社チームアルマダの口座へ着金したことを確認した同日に実行する。", "年利2.7%固定。元金は2026年11月から毎月150,000円、2027年12月の最終回50,000円で完済する。", "唯一株主は肥塚恭子（100%）。株式会社チームアルマダは取締役会非設置・監査役非設置。", "承認済み立替経費は役員借入金、今回の貸付は役員貸付金として別残高で管理し、弁済期到来分のみ月次で相殺する。"]: bullet(x)
heading("2. 株主決定書")
para("唯一株主である肥塚恭子は、取締役である山地正洋との直接取引について、重要な事実の開示を受け、次のとおり事前承認する。")
table(["項目", "決定内容"], [["元本", "2,000,000円"], ["実行条件", "PayPay銀行の新規500万円融資の株式会社チームアルマダ口座着金確認後、同日に振込"], ["利率", "年2.7%（365日の日割計算）"], ["返済", "2026年11月末日から毎月末日に元金150,000円。2027年12月末日に残元金50,000円"], ["相殺", "承認済み立替経費に係る株式会社チームアルマダの山地正洋への債務と、弁済期到来した本貸付の元金・利息を月次相殺計算書により対当額で相殺できる"], ["変更・免除", "利率、期限その他の条件変更または債務免除は別途の株主決定を要する"]], [1.45, 5.05])
para("決定日: ____年__月__日", after=12); para("唯一株主　肥塚恭子　署名: ____________________", after=18)
heading("3. 金銭消費貸借契約書")
para("株式会社チームアルマダ（以下「甲」という。）と、山地正洋（以下「乙」という。）は、次のとおり金銭消費貸借契約を締結する。", after=8)
for title, body in [("第1条（貸付）", "甲は乙に対し、2026年__月__日、2,000,000円を貸し付ける。甲はPayPay銀行の新規500万円融資の入金を確認後、乙指定口座へ振り込む。"), ("第2条（利息）", "利率は年2.7%とし、365日の日割計算で、各返済日に元金とは別に支払う。"), ("第3条（元金返済）", "乙は2026年11月末日から2027年11月末日まで毎月末日に150,000円、2027年12月末日に50,000円を返済する。"), ("第4条（立替経費との相殺）", "承認済み立替経費に係る甲の乙に対する債務と、乙の弁済期到来債務は、月次相殺計算書により対当額で相殺できる。差額のみを銀行振込する。"), ("第5条（期限前返済）", "乙は甲に通知のうえ、元本の全部または一部を期限前に返済できる。"), ("第6条（期限の利益の喪失）", "乙が返済を30日以上遅滞したとき、甲は書面通知により残元利金の一括返済を請求できる。"), ("第7条（変更）", "本契約の変更または債務免除は、甲の株主決定を経た書面合意がある場合に限る。")]:
    para(title, True, after=2); para(body, after=5)
para("契約日: ____年__月__日", after=12); para("甲　株式会社チームアルマダ", after=3); para("　　代表取締役　山地正洋　署名: ____________________", after=6); para("乙　山地正洋　署名: ____________________", after=18)
heading("4. 返済・相殺予定")
rows=[]
for i in range(14):
    y,m=(2026,11+i) if i<2 else (2027,i-1)
    if m>12: y,m=2027,m-12
    principal=50000 if (y,m)==(2027,12) else 150000
    rows.append([f"{y}-{m:02d}", f"{principal:,}円", "元本残高の日割", "承認済み立替額との対当額", "相殺後の差額のみ振込"])
table(["返済月","元金","利息","相殺","実行"],rows,[.8,1.05,1.3,1.9,1.45])
heading("5. freee是正・月次照合")
for x in ["既存の月200,000円の定例行は、貸付返済ではなく、山地正洋立替精算の予定へ再分類する。", "既存の山地正洋への振込は、領収書・承認状況ごとに費用と役員借入金へ振り分ける。", "新規貸付実行日は「役員貸付金 / 普通預金」。元金相殺は「役員借入金 / 役員貸付金」、利息相殺は「役員借入金 / 受取利息」。", "月次で、立替台帳、相殺計算書、freee、銀行明細を照合して残高を確定する。"]: bullet(x)
doc.save(OUT)
print(OUT)
