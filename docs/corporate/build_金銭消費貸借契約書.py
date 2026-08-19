from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = "/Users/masa/projects/AMD/amd-os/docs/corporate/20260816_金銭消費貸借契約書.docx"
doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Inches(0.9)
section.left_margin = section.right_margin = Inches(0.95)

for style_name, size in [("Normal", 10.5), ("Heading 1", 15)]:
    style = doc.styles[style_name]
    style.font.name = "Hiragino Sans"
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:ascii"), "Hiragino Sans")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Hiragino Sans")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans")

def p(text="", bold=False, align=None, after=7, size=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = 1.25
    if align:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Hiragino Sans"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Hiragino Sans")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Hiragino Sans")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans")
    if size:
        run.font.size = Pt(size)
    return para

def article(title, body):
    p(title, bold=True, after=3)
    p(body, after=8)

p("金銭消費貸借契約書", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=20, after=20)
p("株式会社チームアルマダ（以下「甲」という。）と山地正洋（以下「乙」という。）は、次のとおり金銭消費貸借契約（以下「本契約」という。）を締結する。", after=14)

article("第1条（貸付）", "甲は乙に対し、金2,000,000円を貸し付け、乙はこれを借り受ける。甲は、PayPay銀行の新規融資金5,000,000円が甲名義の口座へ着金したことを確認した後、2026年__月__日、乙指定口座へ振り込む方法により貸付金を交付する。")
article("第2条（利息）", "本貸付の利率は年2.7%とする。利息は、貸付実行日から各返済日までの実日数について、1年を365日として日割計算し、元金とは別に各返済日に支払う。")
article("第3条（元金の返済）", "乙は甲に対し、2026年11月末日から2027年11月末日まで、毎月末日に150,000円ずつ元金を返済し、2027年12月末日に残元金50,000円を返済する。")
article("第4条（立替経費債務との相殺）", "甲が乙に対して負担する、甲の業務に係る承認済み立替経費の精算債務と、乙が甲に対して負担する本契約に基づく弁済期到来済みの元金および利息債務は、月次相殺計算書により対当額で相殺することができる。この場合、相殺後の差額のみを銀行振込により精算する。")
article("第5条（期限前返済）", "乙は甲に対し事前に通知することにより、元本の全部または一部を期限前に返済することができる。")
article("第6条（期限の利益の喪失）", "乙が本契約に基づく支払を30日以上遅滞したときは、甲は書面による通知をもって、乙に対し残元利金の一括返済を請求することができる。")
article("第7条（変更）", "本契約の変更、利率の変更、返済期限の延長または債務免除は、甲の株主決定を経た書面による合意がある場合に限り効力を生じる。")
article("第8条（協議）", "本契約に定めのない事項または本契約の解釈に疑義が生じた事項は、甲乙誠実に協議して定める。")

p("本契約締結の証として、本書2通を作成し、甲乙記名押印または電子署名のうえ、各1通を保有する。", after=16)
p("契約日　2026年__月__日", after=16)
p("甲　茨城県つくば市吾妻一丁目10番1号", after=2)
p("　　株式会社チームアルマダ", after=2)
p("　　代表取締役　山地正洋　____________________", after=14)
p("乙　山地正洋　____________________", after=4)

doc.save(OUT)
print(OUT)
